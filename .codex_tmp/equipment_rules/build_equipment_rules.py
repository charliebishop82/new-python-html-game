from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SOURCE = Path(r"C:\Users\charl\OneDrive\Desktop\bbsgame\game\docs\EQUIPMENT_BALANCE_RULES.md")
OUTPUT = Path(r"C:\Users\charl\OneDrive\Desktop\bbsgame\game\docs\Movie_Multiverse_Equipment_Balance_Rules.docx")

NAVY = "20364F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "202124"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


INLINE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def add_inline(paragraph, text, *, color=BLACK, size=11, bold=False):
    parts = INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=10, color="7A3E00", name="Consolas")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, bold=bold, color=color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_table(doc, rows):
    cols = len(rows[0])
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 4:
        widths = [1800, 2520, 1800, 3240]
    else:
        widths = [9360 // cols] * cols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, size=9.25, bold=(r_idx == 0), color=(WHITE if r_idx == 0 else BLACK))
            if r_idx == 0:
                set_cell_shading(cell, BLUE)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    hp = section.header.paragraphs[0]
    hp.text = "MOVIE MULTIVERSE  |  EQUIPMENT DESIGN STANDARD"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        set_font(run, size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            add_inline(p, line[2:], size=24, bold=True, color=NAVY)
            sub = doc.add_paragraph("Equipment power budgets, stacking rules, and import review thresholds")
            sub.paragraph_format.space_after = Pt(16)
            for run in sub.runs:
                set_font(run, size=12, italic=True, color=MUTED)
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("| "):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    rows.append(cells)
                index += 1
            add_table(doc, rows)
            continue
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        else:
            p = doc.add_paragraph()
            if line.startswith("**Status:**") or line.startswith("**Version:**") or line.startswith("**Applies to:**"):
                p.paragraph_format.space_after = Pt(2)
            add_inline(p, line.replace("  ", ""))
        index += 1

    core = doc.core_properties
    core.title = "Movie Multiverse Equipment Balance Rules"
    core.subject = "Authoritative equipment power budgets and validation rules"
    core.author = "Movie Multiverse Project"
    core.keywords = "equipment, balance, weapons, armor, specials, world bosses"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
